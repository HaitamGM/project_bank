import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report(output_path='docs/Rapport_Avancement_BankIA.docx'):
    print("Generating Word document...")
    doc = Document()

    # Title
    title = doc.add_heading('Note d\'Avancement - Projet BankIA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    doc.add_paragraph(
        "Ce document a pour but de présenter l'état d'avancement du projet BankIA, "
        "en mettant l'accent sur la méthodologie de gestion de projet adoptée, "
        "la phase de conception architecturale (UML), ainsi que la stratégie "
        "d'évaluation des performances et des modèles d'Intelligence Artificielle intégrés."
    )

    # 1. Methodology
    doc.add_heading('1. Méthodologie de Projet : Le Cycle en V', level=1)
    doc.add_paragraph(
        "Pour assurer une structure rigoureuse au développement de notre plateforme "
        "bancaire augmentée par l'IA, nous avons opté pour la méthodologie du Cycle en V. "
        "Ce choix garantit un alignement constant entre les besoins métiers (banque) "
        "et les implémentations techniques."
    )
    doc.add_paragraph(
        "Notre implémentation suit les phases suivantes :\n"
        "- Analyse des besoins : Identification des contraintes de conformité et du besoin "
        "d'un assistant vocal intelligent.\n"
        "- Spécifications et Conception (Phase descendante) : Architecture globale (React, FastAPI, WebSockets) "
        "et modélisation UML.\n"
        "- Implémentation : Développement du Moteur Multi-Agents et du système RAG.\n"
        "- Tests et Évaluation (Phase ascendante) : Validation des modèles prédictifs, de la "
        "génération LLM et des performances temps réel."
    )

    # 2. Conception
    doc.add_heading('2. Phase de Conception (Modélisation UML)', level=1)
    doc.add_paragraph(
        "Afin de modéliser les interactions entre les utilisateurs, le système, "
        "et les différents agents d'IA, nous avons élaboré les diagrammes suivants :"
    )

    doc.add_heading('2.1 Diagramme des Cas d\'Utilisation (Système Global)', level=2)
    doc.add_paragraph(
        "Ce diagramme illustre les interactions principales entre le Client (acteur externe), "
        "l'Assistant Vocal, et les systèmes bancaires internes."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_text("[ Insérez ici la capture d'écran du Diagramme des Cas d'Utilisation ]")
    r.font.italic = True
    r.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_heading('2.2 Diagramme de Séquence (Processus : Demande de Crédit)', level=2)
    doc.add_paragraph(
        "Ce diagramme détaille le flux d'orchestration Multi-Agents lors d'une demande "
        "de crédit. Il met en évidence l'appel aux agents de Solvabilité (Scoring) "
        "et à l'agent de Conformité (RAG ChromaDB) de manière séquentielle."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_text("[ Insérez ici la capture d'écran du Diagramme de Séquence ]")
    r.font.italic = True

    doc.add_heading('2.3 Diagramme de Classes', level=2)
    doc.add_paragraph(
        "Ce diagramme représente la structure des données du projet, notamment les entités "
        "Client, DemandeCredit, et RapportExplicabilite."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_text("[ Insérez ici la capture d'écran du Diagramme de Classes ]")
    r.font.italic = True

    # 3. Evaluation and Metrics
    doc.add_heading('3. Évaluation et Mesure des Performances', level=1)
    doc.add_paragraph(
        "L'évaluation de notre architecture BankIA est divisée en trois axes critiques "
        "pour valider à la fois l'exactitude mathématique et la sécurité juridique des décisions."
    )

    doc.add_heading('3.1 Évaluation du Moteur de Scoring (Machine Learning)', level=2)
    doc.add_paragraph(
        "Nous avons mis en place un pipeline d'évaluation ML sur un jeu de données "
        "synthétiques reflétant les règles prudentielles marocaines. Un modèle (ex: Random Forest) "
        "est évalué pour sa capacité à distinguer les dossiers risqués des dossiers solvables."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_text("[ Insérez ici la capture d'écran de la Matrice de Confusion (Confusion Matrix) ]\n\n")
    r.add_text("[ Insérez ici la capture d'écran de la Courbe ROC-AUC ]")
    r.font.italic = True

    doc.add_heading('3.2 Évaluation du Système RAG (Conformité Légale)', level=2)
    doc.add_paragraph(
        "Pour s'assurer que l'IA ne génère pas de fausses lois (Hallucinations), "
        "le système RAG est évalué via des métriques dédiées (inspirées de RAGAS) :\n"
        "- Précision du Contexte (Context Precision) : Pertinence des documents extraits de ChromaDB.\n"
        "- Fidélité (Faithfulness) : Garantie que le rapport final est strictement basé sur le contexte récupéré."
    )

    doc.add_heading('3.3 Performances de l\'Architecture Temps Réel', level=2)
    doc.add_paragraph(
        "Pour assurer une expérience fluide de l'assistant vocal, la latence est monitorée :\n"
        "- Temps d'exécution de la pipeline Multi-Agents (Orchestrateur).\n"
        "- Latence de bout-en-bout (End-to-End latency) via WebSockets (inférieure à 2 secondes)."
    )

    # Conclusion
    doc.add_heading('4. Prochaines Étapes', level=1)
    doc.add_paragraph(
        "Les prochaines étapes consisteront à finaliser l'intégration des métriques, "
        "à affiner l'expérience utilisateur sur le Frontend React, et à consolider "
        "le rapport final de stage."
    )

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Document successfully generated at: {output_path}")

if __name__ == '__main__':
    create_report()
